from datetime import date
from io import BytesIO
import base64
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm


PLACEHOLDERS = {
    "{{nomeContratante}}": "nome_contratante",
    "{{cpfContratante}}": "cpf_contratante",
    "{{rgContratante}}": "rg_contratante",
    "{{cidadeContratante}}": "cidade_contratante",
    "{{enderecoContratante}}": "endereco_contratante",
    "{{numero}}": "numero",
    "{{complemento}}": "complemento",
    "{{bairroContratante}}": "bairro_contratante",
    "{{cepContratante}}": "cep_contratante",
    "{{telefone}}": "telefone",
    "{{email}}": "email",
    "{{nomeContratada}}": "nome_contratada",
    "{{cpfCnpjContratada}}": "cpf_cnpj_contratada",
    "{{cro}}": "cro",
    "{{enderecoClinica}}": "endereco_clinica",
    "{{nomePaciente}}": "nome_paciente",
    "{{procedimentos}}": "procedimentos",
    "{{valorTotal}}": "valor_total",
    "{{formaPagamento}}": "forma_pagamento",
    "{{parcelas}}": "parcelas",
    "{{vencimentos}}": "vencimentos",
    "{{valorAvaliacao}}": "valor_avaliacao",
    "{{limiteAtraso}}": "limite_atraso",
    "{{cidadeClinica}}": "cidade_clinica",
}


def _clean(value):
    return " ".join((value or "").strip().split())


def _replace_in_paragraph(paragraph, replacements):
    for token, value in replacements.items():
        while token in paragraph.text:
            combined = "".join(run.text for run in paragraph.runs)
            start = combined.index(token)
            end = start + len(token)
            cursor = 0
            touched = []
            for run in paragraph.runs:
                run_start, run_end = cursor, cursor + len(run.text)
                if run_end > start and run_start < end:
                    touched.append((run, max(0, start - run_start), min(len(run.text), end - run_start)))
                cursor = run_end
            if not touched:
                break
            first, first_start, first_end = touched[0]
            first.text = first.text[:first_start] + value + first.text[first_end:]
            for run, local_start, local_end in touched[1:]:
                run.text = run.text[:local_start] + run.text[local_end:]


def _all_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def _insert_logo(document, image_source):
    image_source.seek(0)
    for index, section in enumerate(document.sections):
        if index and section.header.is_linked_to_previous:
            continue
        paragraph = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.add_run().add_picture(image_source, width=Cm(2.7))
        image_source.seek(0)


def _insert_patient_signature(document, data_url):
    if not data_url or not data_url.startswith("data:image/png;base64,"):
        return
    try:
        raw = base64.b64decode(data_url.split(",", 1)[1], validate=True)
    except Exception as exc:
        raise ValueError("A assinatura manuscrita recebida é inválida.") from exc
    if len(raw) > 2_000_000:
        raise ValueError("A assinatura manuscrita ultrapassa o tamanho permitido.")
    for paragraph in document.paragraphs:
        if "CONTRATANTE / PACIENTE" in paragraph.text:
            signature = paragraph.insert_paragraph_before()
            signature.alignment = WD_ALIGN_PARAGRAPH.CENTER
            signature.paragraph_format.space_after = 0
            signature.add_run().add_picture(BytesIO(raw), width=Cm(4.2))
            return


def generate_contract_docx(form, template_source, image_source=None):
    if hasattr(template_source, "seek"):
        template_source.seek(0)
    try:
        document = Document(template_source)
    except Exception as exc:
        raise ValueError("Não foi possível abrir o contrato-base. Use um arquivo Word .docx válido.") from exc

    replacements = {token: _clean(form.get(field)) for token, field in PLACEHOLDERS.items()}
    replacements["____ de __________________ de ______"] = _clean(form.get("data_contrato"))
    for paragraph in _all_paragraphs(document):
        _replace_in_paragraph(paragraph, replacements)

    _insert_patient_signature(document, form.get("assinatura_paciente"))

    remaining = sorted({token for paragraph in _all_paragraphs(document) for token in PLACEHOLDERS if token in paragraph.text})
    if remaining:
        raise ValueError("Alguns campos do modelo não puderam ser preenchidos: " + ", ".join(remaining))

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


def default_contract_form():
    hoje = date.today()
    meses = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
    return {
        "nome_contratante": "",
        "cpf_contratante": "",
        "rg_contratante": "",
        "cidade_contratante": "",
        "endereco_contratante": "",
        "numero": "",
        "complemento": "Não se aplica",
        "bairro_contratante": "",
        "cep_contratante": "",
        "telefone": "",
        "email": "",
        "nome_contratada": "Consultório Odontológico Dr. Ângelo G. Martinez",
        "cpf_cnpj_contratada": "",
        "cro": "",
        "endereco_clinica": "",
        "nome_paciente": "",
        "procedimentos": "",
        "valor_total": "",
        "forma_pagamento": "",
        "parcelas": "À vista",
        "vencimentos": "Conforme orçamento/financeiro",
        "valor_avaliacao": "",
        "limite_atraso": "15",
        "cidade_clinica": "",
        "data_contrato": f"{hoje.day:02d} de {meses[hoje.month - 1]} de {hoje.year}",
    }
